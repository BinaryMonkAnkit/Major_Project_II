import os
import fitz  # PyMuPDF for PDFs
import pandas as pd
import uuid
import docx
from docx import Document
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
import io
from django.conf import settings
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.html import partition_html
from django.utils.text import slugify

MEDIA_ROOT = settings.MEDIA_ROOT
MEDIA_URL = settings.MEDIA_URL  # Ensure this is defined in settings.py

class DocumentProcessingService:
    def __init__(self, document_path, session_id):
        self.document_path = document_path
        self.session_id = session_id
        self.temp_dir = os.path.join(MEDIA_ROOT, "temp", session_id)
        os.makedirs(self.temp_dir, exist_ok=True)

    def extract_text(self):
        """
        Extracts text from different document types and saves it as a temporary knowledge base.
        """
        text = ""
        file_extension = os.path.splitext(self.document_path)[1].lower()

        try:
            if file_extension == ".pdf":
                text = self.extract_text_from_pdf()
            elif file_extension == ".docx" or file_extension == ".doc":
                text = self.extract_text_from_docx()
            elif file_extension == ".html":
                text = self.extract_text_from_html()
            elif file_extension == ".txt":
                text = self.extract_text_from_txt()
            else:
                return None

            text_file_path = os.path.join(self.temp_dir, f"{self.session_id}_text.txt")
            with open(text_file_path, "w", encoding="utf-8") as f:
                f.write(text)
            return text_file_path

        except Exception as e:
            print(f"Error extracting text: {e}")
            return None

    def extract_text_from_pdf(self):
        """
        Extracts text from a PDF file (normal or scanned).
        """
        text = ""
        try:
            with fitz.open(self.document_path) as doc:
                for page in doc:
                    text += page.get_text("text") + "\n"

            if not text.strip():  # If no text was extracted (possibly a scanned PDF)
                text = self.extract_text_from_scanned_pdf()

            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_scanned_pdf(self):
        """
        Uses OCR (Tesseract) to extract text from scanned PDFs (image-based).
        """
        text = ""
        try:
            with fitz.open(self.document_path) as doc:
                for page_num in range(doc.page_count):
                    pix = doc[page_num].get_pixmap()
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    text += pytesseract.image_to_string(img) + "\n"
        except Exception as e:
            print(f"Error extracting text from scanned PDF: {e}")
        return text

    def extract_text_from_docx(self):
        """
        Extracts text from DOCX files.
        """
        text = ""
        try:
            doc = Document(self.document_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
        return text

    def extract_text_from_html(self):
        """
        Extracts text from HTML files.
        """
        text = ""
        try:
            with open(self.document_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                text = soup.get_text(separator="\n")
        except Exception as e:
            print(f"Error extracting text from HTML: {e}")
        return text

    def extract_text_from_txt(self):
        """
        Extracts text from TXT files.
        """
        text = ""
        try:
            with open(self.document_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
        return text

    def extract_tables(self):
        """
        Extracts tables and saves them as CSV files in the temp directory.
        Supports PDF, HTML, DOCX, DOC, and TXT formats.
        """
        file_extension = os.path.splitext(self.document_path)[1].lower()

        try:
            if file_extension == ".pdf":
                return self.extract_tables_from_pdf()
            elif file_extension == ".html":
                return self.extract_tables_from_html()
            elif file_extension in [".docx", ".doc"]:
                return self.extract_tables_from_docx()
            elif file_extension == ".txt":
                return self.extract_tables_from_txt()
            else:
                return []
        except Exception as e:
            print(f"Error extracting tables: {e}")
            return []

    def extract_tables_from_pdf(self):
        """
        Extracts tables from a PDF file and saves them as CSV.
        """
        tables_data = []
        try:
            #Recognize the table structure in the PDF
            elements = partition_pdf(self.document_path, infer_table_structure=True, strategy='hi_res')
            tables = [el for el in elements if el.category == "Table"]

            for table in tables:
                html = table.metadata.text_as_html
                dfs = pd.read_html(html)
                if dfs:
                    tables_data.append(dfs[0])

            return self._save_tables_as_csv(tables_data, "pdf")
        except Exception as e:
            print(f"Error extracting tables from PDF: {e}")
            return []

    def extract_tables_from_html(self):
        """
        Extracts tables from an HTML file and saves them as CSV.
        """
        try:
            dfs = pd.read_html(self.document_path)
            return self._save_tables_as_csv(dfs, "html")
        except Exception as e:
            print(f"Error extracting tables from HTML: {e}")
            return []

    def extract_tables_from_docx(self):
        """
        Extracts tables from DOCX files.
        """
        tables_data = []
        try:
            doc = Document(self.document_path)
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                df = pd.DataFrame(rows)
                tables_data.append(df)

            return self._save_tables_as_csv(tables_data, "docx")
        except Exception as e:
            print(f"Error extracting tables from DOCX/DOC: {e}")
            return []

    def extract_tables_from_txt(self):
        """
        Extracts structured data from TXT files if applicable.
        """
        tables_data = []
        try:
            with open(self.document_path, "r", encoding="utf-8") as f:
                lines = [line.strip().split() for line in f if line.strip()]
            
            if lines:
                df = pd.DataFrame(lines)
                tables_data.append(df)

            return self._save_tables_as_csv(tables_data, "txt")
        except Exception as e:
            print(f"Error processing TXT file: {e}")
            return []

    def _save_tables_as_csv(self, tables_data, file_type):
        """
        Helper function to save extracted tables as CSV files.
        """
        timestamp = uuid.uuid4().hex
        csv_files = []

        for i, df in enumerate(tables_data):
            file_name = f"{self.session_id}_table_{file_type}_{i}_{timestamp}.csv"
            csv_file_path = os.path.join(self.temp_dir, file_name)
            df.to_csv(csv_file_path, index=False)

            # Generate the full URL
            file_url = f"{settings.MEDIA_URL}temp/{self.session_id}/{file_name}"
            csv_files.append({
                "name": file_name,
                "path": file_url,  # Ensure full URL is returned
                "type": "csv"
            })

        return csv_files

    # def cleanup(self):
    #     """
    #     Deletes all extracted text and tables for the session after it's done.
    #     """
    #     try:
    #         if os.path.exists(self.temp_dir):
    #             for file in os.listdir(self.temp_dir):
    #                 os.remove(os.path.join(self.temp_dir, file))
    #             os.rmdir(self.temp_dir)
    #     except Exception as e:
    #         print(f"Error during cleanup: {e}")
